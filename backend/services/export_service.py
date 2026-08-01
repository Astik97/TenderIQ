"""
=========================================================
TenderIQ Export Service
---------------------------------------------------------

1. JSON Export

2. TXT Export

3. DOCX Export

4. PDF Export

=========================================================
"""

import json

from io import BytesIO

from docx import Document

from reportlab.platypus import (SimpleDocTemplate,Paragraph)

from reportlab.lib.styles import getSampleStyleSheet

# =========================================================
# JSON Export
# =========================================================

def export_json(result):

    return json.dumps(result,indent=4,default=str)

# =========================================================
# TXT Export
# =========================================================

def export_txt(result):

    lines = []

    lines.append("TenderIQ AI Comparison Report")

    lines.append("=" * 50)

    lines.append("")

    lines.append(f"Overall Similarity : {result['similarity']}%")

    lines.append(f"Weighted Similarity : {result['weighted_similarity']}%")

    lines.append("")

    lines.append("Clause Comparison")

    lines.append("-" * 40)

    for clause in result["clause_results"]:

        lines.append("")

        lines.append(f"Similarity : {clause['similarity']}%")

        lines.append(f"Risk : {clause['risk']['level']}")

        lines.append(f"Clause : {clause['clause']}")

    return "\n".join(lines)

# =========================================================
# DOCX Export
# =========================================================

def export_docx(result):

    document = Document()

    document.add_heading("TenderIQ AI Comparison Report",level=1)

    document.add_heading("Summary",level=2)

    document.add_paragraph(f"Overall Similarity : {result['similarity']}%")

    document.add_paragraph(f"Weighted Similarity : {result['weighted_similarity']}%")

    document.add_paragraph(f"Total Clauses : {result['total_clauses']}")

    document.add_heading("Clause Results",level=2)

    for clause in result["clause_results"]:

        document.add_heading(clause["clause"][:80],level=3)

        document.add_paragraph(f"Similarity : {clause['similarity']}%")

        document.add_paragraph(f"Risk : {clause['risk']['level']}")

        document.add_paragraph(f"Recommendation : {clause['risk']['recommendation']}")

    buffer = BytesIO()

    document.save(buffer)

    buffer.seek(0)

    return buffer

# =========================================================
# PDF Export
# =========================================================

def export_pdf(result):

    buffer = BytesIO()

    document = SimpleDocTemplate(buffer)

    styles = getSampleStyleSheet()

    story = []

    story.append(Paragraph(
        "TenderIQ AI Comparison Report",
        styles["Heading1"]
        ))

    story.append(Paragraph(
        f"Overall Similarity : {result['similarity']}%",
        styles["Normal"]
        ))

    story.append(Paragraph(
        f"Weighted Similarity : {result['weighted_similarity']}%",
        styles["Normal"]
        ))

    story.append(Paragraph("",styles["Normal"]))

    for clause in result["clause_results"]:

        story.append(Paragraph(
            f"<b>Clause</b>: {clause['clause']}",
            styles["BodyText"]
            ))

        story.append(Paragraph(
            f"Similarity : {clause['similarity']}%",
            styles["BodyText"]
        ))

        story.append(Paragraph(
            f"Risk : {clause['risk']['level']}",
            styles["BodyText"]
            ))

        story.append(Paragraph("",styles["BodyText"]))

    document.build(story)

    buffer.seek(0)

    return buffer

# =========================================================
# Export Dispatcher
# =========================================================

def export_report(result, file_type):

    file_type = file_type.lower()

    if file_type == "json":

        return export_json(result)

    elif file_type == "txt":

        return export_txt(result)

    elif file_type == "docx":

        return export_docx(result)

    elif file_type == "pdf":

        return export_pdf(result)

    raise ValueError("Unsupported export format.")